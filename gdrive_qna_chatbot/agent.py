from datetime import datetime
import hashlib
import json
import logging
import os

# Vector database
import chromadb  # Using ChromaDB as vector database
import docx  # python-docx for Word documents
# Document processing
import fitz  # PyMuPDF for PDF processing
# Gemini integration
import google.generativeai as genai
import pandas as pd  # pandas for Excel files
# Embeddings generation
from sentence_transformers import SentenceTransformer  # Local embedding model


# Optional OCR for images - will be checked at runtime
# pip install pytesseract pillow
# Also requires Tesseract OCR to be installed on the system

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class LocalFileProcessor:
    """Handles local file discovery and content extraction"""

    def __init__(self, docs_directory):
        """Initialize with the path to the documents directory"""
        self.docs_directory = docs_directory

    def list_files(self, recursive=True):
        """List all files in the specified directory, optionally recursively"""
        all_files = []
        supported_extensions = [
            '.pdf', '.docx', '.doc', '.gdoc',  # Document formats
            '.xlsx', '.xls', '.gsheet',         # Spreadsheet formats
            '.jpg', '.jpeg', '.png'             # Image formats
        ]

        if recursive:
            for root, _, files in os.walk(self.docs_directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    extension = os.path.splitext(file)[1].lower()

                    # Only process supported file types
                    if extension in supported_extensions:
                        all_files.append({
                            'id': self._generate_file_id(file_path),
                            'name': file,
                            'path': file_path,
                            'mimeType': self._guess_mime_type(file),
                            'modifiedTime': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
                        })
        else:
            for file in os.listdir(self.docs_directory):
                file_path = os.path.join(self.docs_directory, file)
                extension = os.path.splitext(file)[1].lower()

                # Only process supported file types
                if os.path.isfile(file_path) and extension in supported_extensions:
                    all_files.append({
                        'id': self._generate_file_id(file_path),
                        'name': file,
                        'path': file_path,
                        'mimeType': self._guess_mime_type(file),
                        'modifiedTime': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
                    })

        # Log summary of found files
        file_types = {}
        for file_info in all_files:
            ext = os.path.splitext(file_info['name'])[1].lower()
            file_types[ext] = file_types.get(ext, 0) + 1

        logger.info(f"Found {len(all_files)} supported files:")
        for ext, count in file_types.items():
            logger.info(f"  - {ext}: {count} files")

        return all_files

    def _generate_file_id(self, file_path):
        """Generate a unique ID for a file based on its path"""
        return hashlib.md5(file_path.encode()).hexdigest()

    def _guess_mime_type(self, filename):
        """Guess MIME type based on file extension"""
        extension = os.path.splitext(filename)[1].lower()

        mime_types = {
            # Document formats
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.gdoc': 'application/vnd.google-apps.document',

            # Spreadsheet formats
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.gsheet': 'application/vnd.google-apps.spreadsheet',

            # Image formats
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png'
        }

        return mime_types.get(extension, 'application/octet-stream')

class DocumentProcessor:
    """Processes documents into text chunks for embedding"""

    def extract_text(self, file_path, mime_type):
        """Extract text from various file formats"""
        try:
            # PDF documents
            if "pdf" in mime_type:
                return self._extract_from_pdf(file_path)

            # Word documents (DOCX, DOC) and Google Docs
            elif "word" in mime_type or "document" in mime_type or file_path.endswith((".docx", ".doc", ".gdoc")):
                return self._extract_from_docx(file_path)

            # Excel spreadsheets and Google Sheets
            elif "spreadsheet" in mime_type or "excel" in mime_type or file_path.endswith((".xlsx", ".xls", ".gsheet")):
                return self._extract_from_excel(file_path)

            # Image files
            elif "image" in mime_type or file_path.endswith((".jpg", ".jpeg", ".png")):
                return self._extract_from_image(file_path)

            else:
                logger.warning(f"Unsupported mime type: {mime_type} for {file_path}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _extract_from_pdf(self, file_path):
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
        return text

    def _extract_from_docx(self, file_path):
        """Extract text from Word documents and Google Docs"""
        try:
            # For Google Docs (.gdoc files)
            if file_path.endswith('.gdoc'):
                try:
                    # .gdoc files are actually JSON files with a reference URL
                    with open(file_path, 'r', encoding='utf-8') as f:
                        gdoc_data = json.load(f)

                    # Extract any available text content from the JSON
                    if 'content' in gdoc_data:
                        return gdoc_data['content']
                    else:
                        logger.warning(f"Google Doc {file_path} does not contain direct content")
                        return f"[Google Doc: {os.path.basename(file_path)}] - Content could not be extracted directly"
                except Exception as e:
                    logger.error(f"Error processing Google Doc: {str(e)}")
                    return ""

            # For standard Word documents
            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs]

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)

            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting Word document text: {str(e)}")
            return ""

    def _extract_from_excel(self, file_path):
        """Extract text from Excel and Google Sheets"""
        try:
            # For Google Sheets (.gsheet files)
            if file_path.endswith('.gsheet'):
                try:
                    # .gsheet files are JSON files with reference URLs
                    with open(file_path, 'r', encoding='utf-8') as f:
                        gsheet_data = json.load(f)

                    # Extract any available text content from the JSON
                    if 'content' in gsheet_data:
                        return gsheet_data['content']
                    else:
                        logger.warning(f"Google Sheet {file_path} does not contain direct content")
                        return f"[Google Sheet: {os.path.basename(file_path)}] - Content could not be extracted directly"
                except Exception as e:
                    logger.error(f"Error processing Google Sheet: {str(e)}")
                    return ""

            # For Excel files
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets

            # Combine all sheets into a single text representation
            all_sheets = []
            for sheet_name, sheet_df in df.items():
                all_sheets.append(f"Sheet: {sheet_name}")
                all_sheets.append(sheet_df.to_string(index=False))

            return "\n\n".join(all_sheets)
        except Exception as e:
            logger.error(f"Error extracting spreadsheet text: {str(e)}")
            return ""

    def _extract_from_image(self, file_path):
        """Extract text from images using OCR if available"""
        try:
            # Check if pytesseract is installed
            import importlib
            if importlib.util.find_spec("pytesseract") is not None:
                try:
                    import pytesseract
                    from PIL import Image

                    # Open image and convert to RGB (in case it's RGBA or other format)
                    image = Image.open(file_path).convert('RGB')

                    # Use pytesseract to extract text
                    text = pytesseract.image_to_string(image)

                    if text.strip():
                        return text
                    else:
                        return f"[Image: {os.path.basename(file_path)}] - No text detected"
                except Exception as ocr_e:
                    logger.error(f"OCR error for {file_path}: {str(ocr_e)}")
                    return f"[Image: {os.path.basename(file_path)}] - OCR processing failed"
            else:
                logger.warning("pytesseract not installed. Install with 'pip install pytesseract' and install Tesseract OCR")
                return f"[Image: {os.path.basename(file_path)}] - OCR not available"
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            return f"[Image: {os.path.basename(file_path)}]"

    def chunk_text(self, text, chunk_size=1000, overlap=200):
        """Split text into overlapping chunks"""
        chunks = []
        if not text:
            return chunks

        words = text.split()

        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk:
                chunks.append(chunk)

        return chunks

class VectorStore:
    """Manages document embeddings in a vector database"""

    def __init__(self, db_path="./vectordb"):
        self.client = chromadb.PersistentClient(path=db_path)
        self.collection = self.client.get_or_create_collection("local_documents")
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')  # Small but effective model

    def generate_embeddings(self, texts):
        """Generate embeddings for a list of texts"""
        return self.embedding_model.encode(texts).tolist()

    def store_document(self, doc_id, chunks, file_metadata):
        """Store document chunks with their embeddings"""
        # Generate unique IDs for each chunk
        chunk_ids = [f"{doc_id}-chunk-{i}" for i in range(len(chunks))]

        # Create embeddings
        embeddings = self.generate_embeddings(chunks)

        # Create metadata for each chunk
        metadatas = [{"doc_id": doc_id,
                      "chunk_index": i,
                      "filename": file_metadata["name"],
                      "file_path": file_metadata["path"],
                      "mime_type": file_metadata["mimeType"],
                      "modified_time": file_metadata["modifiedTime"],
                      "total_chunks": len(chunks)}
                     for i in range(len(chunks))]

        # Add to collection
        self.collection.add(
            ids=chunk_ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas
        )

        return chunk_ids

    def delete_document(self, doc_id):
        """Remove all chunks for a specific document"""
        self.collection.delete(where={"doc_id": doc_id})

    def query(self, query_text, n_results=5):
        """Find relevant document chunks for a query"""
        query_embedding = self.embedding_model.encode([query_text]).tolist()
        results = self.collection.query(
            query_embeddings=query_embedding,
            n_results=n_results,
            include=["documents", "metadatas", "distances"]
        )

        return results

class GeminiClient:
    """Interface to Google's Gemini API"""

    def __init__(self, api_key=None, model="gemini-2.0-flash"):
        """Initialize Gemini client

        Args:
            api_key: Google API key for Gemini (can be set via GOOGLE_API_KEY env variable)
            model: Gemini model to use (default: gemini-1.5-pro)
        """
        self.model = model

        # Configure Gemini API with key (from param or environment variable)
        if api_key is None:
            api_key = os.environ.get("GOOGLE_API_KEY")
            if api_key is None:
                logger.error("No API key provided for Gemini. Set GOOGLE_API_KEY environment variable or pass api_key parameter.")
                raise ValueError("Gemini API key is required. Set GOOGLE_API_KEY environment variable or pass api_key parameter.")

        # Configure the Gemini API
        genai.configure(api_key=api_key)

        # Check if the model is available
        try:
            models = [m.name for m in genai.list_models()]
            if self.model not in models:
                available_models = [m for m in models if "gemini" in m]
                if available_models:
                    logger.warning(f"Model '{self.model}' not found. Available Gemini models: {available_models}")
                    # Use the first available Gemini model as fallback
                    self.model = available_models[0]
                    logger.info(f"Using '{self.model}' as fallback")
                else:
                    logger.error("No Gemini models available")
                    raise ValueError("No Gemini models available")
            else:
                logger.info(f"Using Gemini model: {self.model}")
        except Exception as e:
            logger.error(f"Error checking Gemini models: {str(e)}")
            # Continue anyway, might work with the default model

    def generate(self, prompt, system_prompt=None, temperature=0.1):
        """Generate response from Gemini

        Args:
            prompt: The user prompt
            system_prompt: System instructions (optional)
            temperature: Generation temperature (0.0 to 1.0)

        Returns:
            Generated text response
        """
        try:
            # Configure model generation parameters
            generation_config = {
                "temperature": temperature,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 2048,
            }

            # Create model instance
            model = genai.GenerativeModel(
                model_name=self.model,
                generation_config=generation_config
            )

            if system_prompt:
                # Prepend system instructions to the user prompt
                enhanced_prompt = f"{system_prompt}\n\n{prompt}"
                logger.info("Added system instructions to user prompt")
            else:
                enhanced_prompt = prompt

            # Generate response
            response = model.generate_content(enhanced_prompt)

            # Return the text response
            if response.text:
                return response.text
            else:
                logger.error("Gemini returned empty response")
                return "Error: No response generated from Gemini."

        except Exception as e:
            logger.error(f"Error calling Gemini API: {str(e)}")
            return f"Error generating response: {str(e)}"


class IndexBuilder:
    """Builds the document index from local files"""

    def __init__(self, file_processor, doc_processor, vector_store):
        self.file_processor = file_processor
        self.doc_processor = doc_processor
        self.vector_store = vector_store
        self.processed_files = {}  # Track processed files
        self.processed_files_path = "processed_files.json"
        self._load_processed_files()

    def _load_processed_files(self):
        """Load registry of processed files if it exists"""
        if os.path.exists(self.processed_files_path):
            try:
                with open(self.processed_files_path, 'r') as f:
                    self.processed_files = json.load(f)
            except Exception as e:
                logger.error(f"Error loading processed files record: {str(e)}")
                self.processed_files = {}

    def _save_processed_files(self):
        """Save registry of processed files"""
        try:
            with open(self.processed_files_path, 'w') as f:
                json.dump(self.processed_files, f)
        except Exception as e:
            logger.error(f"Error saving processed files record: {str(e)}")

    def build_index(self, force_rebuild=False):
        """Build or update index of all documents"""
        logger.info("Building document index...")

        # Get all files in the documents directory
        files = self.file_processor.list_files(recursive=True)

        indexed_count = 0
        skipped_count = 0

        for file_info in files:
            file_id = file_info['id']
            file_path = file_info['path']
            file_mod_time = file_info['modifiedTime']

            # Skip if already processed and not forcing rebuild
            if not force_rebuild and file_id in self.processed_files:
                if self.processed_files[file_id].get('modifiedTime') == file_mod_time:
                    skipped_count += 1
                    continue

            # Extract text from file
            text = self.doc_processor.extract_text(file_path, file_info['mimeType'])

            # Delete existing document chunks if any
            if file_id in self.processed_files:
                self.vector_store.delete_document(file_id)

            # Process and store if text was extracted
            if text:
                chunks = self.doc_processor.chunk_text(text)
                if chunks:
                    self.vector_store.store_document(file_id, chunks, file_info)

                    # Update processed files registry
                    self.processed_files[file_id] = {
                        "name": file_info["name"],
                        "path": file_path,
                        "mimeType": file_info["mimeType"],
                        "modifiedTime": file_mod_time,
                        "chunk_count": len(chunks),
                        "indexed_at": datetime.now().isoformat()
                    }

                    indexed_count += 1

        self._save_processed_files()
        logger.info(f"Indexing complete. Indexed {indexed_count} files, skipped {skipped_count} files.")
        return indexed_count, skipped_count

class QueryEngine:
    """Handles query processing and RAG implementation"""

    def __init__(self, vector_store, gemini_client):
        self.vector_store = vector_store
        self.gemini_client = gemini_client

    def process_query(self, query, n_results=5):
        """Process a user query using RAG"""
        # Get relevant document chunks
        results = self.vector_store.query(query, n_results=n_results)

        # If no results, return direct LLM response
        if not results or len(results["documents"][0]) == 0:
            return {
                "response": self.gemini_client.generate(query),
                "sources": []
            }

        # Prepare context from retrieved documents
        contexts = []
        sources = []

        for i, (doc, metadata) in enumerate(zip(results["documents"][0], results["metadatas"][0])):
            contexts.append(f"[Document {i+1}]: {doc}")
            sources.append({
                "filename": metadata["filename"],
                "file_path": metadata["file_path"],
                "chunk_index": metadata["chunk_index"],
                "relevance_score": 1 - results["distances"][0][i]  # Convert distance to similarity
            })

        # Create RAG prompt
        rag_prompt = f"""Answer the following question based on the provided context. If the context doesn't contain relevant information, say that you don't have enough information.

Context:
{chr(10).join(contexts)}

Question: {query}

Answer:"""

        # System prompt that instructs the model on how to format responses
        system_prompt = """You are an assistant that answers questions based solely on the provided context.
Cite specific documents when providing information. If the provided context doesn't contain enough information to answer the question,
say that you don't have enough information."""

        # Generate response
        response = self.gemini_client.generate(rag_prompt, system_prompt=system_prompt)
        return {
            "response": response,
            "sources": sources
        }

class ChatInterface:
    """Simple command-line chat interface"""

    def __init__(self, query_engine):
        self.query_engine = query_engine
        self.conversation_history = []

    def start(self):
        """Start the chat interface"""
        print("Local Files ChatBot")
        print("Type 'exit' to quit")

        while True:
            user_input = input("\nYou: ")

            if user_input.lower() in ['exit', 'quit']:
                break

            response_data = self.query_engine.process_query(user_input)
            self.conversation_history.append({"user": user_input, "assistant": response_data})

            print("\nAssistant:", response_data["response"])

            if response_data["sources"]:
                print("\nSources:")
                for i, source in enumerate(response_data["sources"]):
                    print(f"  {i+1}. {source['filename']} (Relevance: {source['relevance_score']:.2f})")

def main():
    """Main application entry point"""
    # Configuration
    docs_directory = "./documents"
    db_path = "./vectordb"

    # Check for Gemini API key from environment
    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        logger.warning("GOOGLE_API_KEY environment variable not set. Please set it before running.")
        print("\nPlease set your Google API key as an environment variable:")
        print("  export GOOGLE_API_KEY='your-api-key'")
        print("Or provide it when prompted.")
        api_key = input("\nEnter your Google Gemini API key: ")

        if not api_key.strip():
            logger.error("No API key provided. Exiting.")
            return

    # Check for image processing dependencies
    try:
        import pytesseract
        from PIL import Image
        logger.info("OCR capabilities available for image processing")
        ocr_available = True
    except ImportError:
        logger.warning("pytesseract and/or PIL not installed. Images will be processed without OCR.")
        logger.warning("Install with: pip install pytesseract pillow")
        logger.warning("And install Tesseract OCR from: https://github.com/tesseract-ocr/tesseract")
        ocr_available = False

    # Initialize components
    file_processor = LocalFileProcessor(docs_directory)
    doc_processor = DocumentProcessor()
    vector_store = VectorStore(db_path=db_path)

    # Initialize Gemini client with API key
    try:
        gemini_client = GeminiClient(api_key=api_key, model="models/gemini-2.0-flash")
    except Exception as e:
        logger.error(f"Failed to initialize Gemini client: {str(e)}")
        return


    # Print supported file types
    logger.info("Processing files with the following extensions:")
    logger.info("  - Documents: .pdf, .docx, .doc, .gdoc")
    logger.info("  - Spreadsheets: .xlsx, .xls, .gsheet")
    logger.info("  - Images: .jpg, .jpeg, .png" + (" (with OCR)" if ocr_available else " (no OCR available)"))

    # Build index
    index_builder = IndexBuilder(file_processor, doc_processor, vector_store)
    index_builder.build_index()

    # Initialize query engine
    query_engine = QueryEngine(vector_store, gemini_client)

    # Start chat interface
    chat_interface = ChatInterface(query_engine)
    chat_interface.start()

if __name__ == "__main__":
    main()
